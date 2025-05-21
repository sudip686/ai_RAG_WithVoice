from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

def load_and_process_docx(docx_path):
    doc = Document(docx_path)
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )
    # Wrap as a LangChain Document
    from langchain.docstore.document import Document as LCDocument
    docs = [LCDocument(page_content=full_text)]
    return splitter.split_documents(docs)